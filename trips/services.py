from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Mm, Inches
from datetime import datetime
import pytils
import zipfile
import io
from functools import reduce
from django.http import HttpResponse
from decimal import Decimal


def insert_event(doc: Document, sight: str, body: str):
    day1_desc = doc.add_paragraph()
    runner = day1_desc.add_run(
        f"{sight.start_at} - {sight.end_at} - {sight.sight.title}"
    )
    runner.bold = False
    runner1 = day1_desc.add_run(body)
    runner1.bold = False
    runner.add_break()


def expense(trip: "Trip"):
    days = trip.days.all()
    total = 0
    for day in days:
        vehicles = day.vehiclestrips_set.all()
        vehicles_price = reduce(
            lambda acc, curr: acc + curr.price * curr.quantity,
            vehicles,
            Decimal("0.00"),
        )
        guides = day.guidetrips_set.all()
        guides_price = reduce(
            lambda acc, curr: acc + curr.price * curr.quantity, guides, Decimal("0.00")
        )
        sights = day.daysight_set.all()
        sights_adult_price = reduce(
            lambda acc, curr: acc + curr.adult_price * curr.adults_quantity,
            sights,
            Decimal("0.00"),
        )
        sights_kid_price = reduce(
            lambda acc, curr: acc + curr.kid_price * curr.kids_quantity,
            sights,
            Decimal("0.00"),
        )
        day_price = vehicles_price + guides_price + sights_adult_price + sights_kid_price
        total += day_price
    return total


def download_program(file_bytes: bytes, filename: str):
    response = HttpResponse(file_bytes,
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",)
    response["Content-Disposition"] = f"attachment; filename={filename}"
    return response


def create_program_docx(trip: "Trip"):
    doc: Document = Document()
    style = doc.styles["Normal"]
    style.font.name = "TimesNewRoman"
    style.font.size = Pt(14)
    style.font.bold = False
    section = doc.sections[0]
    section.left_margin = Mm(15)
    section.right_margin = Mm(10)
    section.top_margin = Mm(15)
    section.bottom_margin = Mm(10)

    header = doc.add_paragraph()
    runner = header.add_run("Программа")
    header.add_run().add_break()
    runner1 = header.add_run(trip.title)
    runner.font.name = runner1.font.name = "Arial"
    runner.font.size = runner1.font.size = Pt(13)
    runner.bold = runner1.bold = True
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.add_run().add_break()

    subheader = doc.add_paragraph()
    subheader.alignment = WD_ALIGN_PARAGRAPH.CENTER
    total = expense(trip)
    runner = subheader.add_run(f"При группе {trip.group()} чел. - {total} руб.")
    runner.bold = True
    runner.add_break()

    days = trip.days.order_by("date").all()

    for i, day in enumerate(days, start=1):
        day_header = doc.add_paragraph()
        day_header.paragraph_format.space_before = Inches(0)
        day_header.paragraph_format.space_after = Inches(0)
        day_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        runner = day_header.add_run(f"{i} день")
        runner.add_break()
        runner.bold = runner1.bold = True
        runner1 = day_header.add_run(
            pytils.dt.ru_strftime("%d %B", inflected=True, date=day.date)
        )
        runner.bold = runner1.bold = True
        runner1.underline = True
        runner1.add_break()

        for sight in day.daysight_set.order_by("start_at").all():
            insert_event(doc, sight, sight.sight.description)

    filename = f'program_{trip.title}_{trip.group()} person_{datetime.today().strftime("%d.%m.%Y")}.docx'
    docx_stream = io.BytesIO()
    doc.save(docx_stream)
    docx_stream.seek(0)

    return (docx_stream.read(), filename)


def archive_program(queryset):
    zip_in_bytes = io.BytesIO()
    zip_filename = f'Program_{datetime.today().strftime("%d.%m.%Y")}.zip'
    for trip in queryset:
        if queryset.count() == 1:
            return queryset.first().download_program()
        with zipfile.ZipFile(zip_in_bytes, mode="a") as archive:
            trips = [trip.generate_program()]
            for file_bytes, filename in trips:
                archive.writestr(filename, file_bytes)
        response = HttpResponse(zip_in_bytes.getvalue(),
                                content_type="application/zip")
        response["Content-Disposition"] = f"attachment; filename={zip_filename}"
    return response
