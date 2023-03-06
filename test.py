from enum import Enum

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


class StylesEnum(Enum):
    COMMENT_STYLE = 'CommentsStyle'


def set_font_style(doc: Document, style: StylesEnum, font_size: int = 14, font: str = 'Times New Roman'):
    font_styles = doc.styles
    font_charstyle = font_styles.add_style(style.value, WD_STYLE_TYPE.CHARACTER)
    font_object = font_charstyle.font
    font_object.size = Pt(font_size)
    font_object.name = font


def insert_event(doc: Document, header: str, body: str):
    day1_desc = doc.add_paragraph()
    runner = day1_desc.add_run(header)
    runner.bold = True
    runner.add_break()
    runner = day1_desc.add_run(body)
    runner.add_break()


def main():
    doc: Document = Document()
    header = doc.add_paragraph()
    header.add_run('Тур в Тюмень').bold = True
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.add_run().add_break()

    subheader = doc.add_paragraph()
    subheader.alignment = WD_ALIGN_PARAGRAPH.CENTER
    runner = subheader.add_run('24-25 февраля')
    runner.bold = True
    runner.add_break()

    desc = doc.add_paragraph()
    runner = desc.add_run('Предполагаемый транспорт: минивэн')
    runner.add_break()

    day1_header = doc.add_paragraph()
    day1_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    runner = day1_header.add_run('24 февраля')
    runner.bold = True
    runner.underline = True
    events = [{'header': '10:00 – 12:00 - Обзорная экскурсия по Тюмени', 'body': '''Гостей ждет увлекательная поездка - прогулка по историческому центру нашего уютного и современного города. Откроем красоту тюменских площадей, историю тюменских улиц и памятников, тайны замысловатых узоров традиционной сибирской резьбы на наличниках старинных домов. Обязательно прогуляемся по Мосту Влюбленных и единственной в России 4-х уровневой набережной. Пройдемся по территории Свято-Троицкого мужского монастыря, уютно расположившегося на территории Затюменки.
# В ходе экскурсии гостей города ждет  приятная встреча с купчихой Валюшей Колокольниковой, которая не только поведает об истории своей семьи, но  и угостит чаем из дымящегося самовара, сибирскими наливками и тюменскими конфетами'''},
              {'header': '10:00 – 12:00 - Обзорная экскурсия по Тюмени', 'body': '''Гостей ждет увлекательная поездка - прогулка по историческому центру нашего уютного и современного города. Откроем красоту тюменских площадей, историю тюменских улиц и памятников, тайны замысловатых узоров традиционной сибирской резьбы на наличниках старинных домов. Обязательно прогуляемся по Мосту Влюбленных и единственной в России 4-х уровневой набережной. Пройдемся по территории Свято-Троицкого мужского монастыря, уютно расположившегося на территории Затюменки.
# # В ходе экскурсии гостей города ждет  приятная встреча с купчихой Валюшей Колокольниковой, которая не только поведает об истории своей семьи, но  и угостит чаем из дымящегося самовара, сибирскими наливками и тюменскими конфетами'''},
              {'header': '10:00 – 12:00 - Обзорная экскурсия по Тюмени', 'body': '''Гостей ждет увлекательная поездка - прогулка по историческому центру нашего уютного и современного города. Откроем красоту тюменских площадей, историю тюменских улиц и памятников, тайны замысловатых узоров традиционной сибирской резьбы на наличниках старинных домов. Обязательно прогуляемся по Мосту Влюбленных и единственной в России 4-х уровневой набережной. Пройдемся по территории Свято-Троицкого мужского монастыря, уютно расположившегося на территории Затюменки.
# # В ходе экскурсии гостей города ждет  приятная встреча с купчихой Валюшей Колокольниковой, которая не только поведает об истории своей семьи, но  и угостит чаем из дымящегося самовара, сибирскими наливками и тюменскими конфетами'''}]

    #     day1_desc = doc.add_paragraph()
    #     runner = day1_desc.add_run('10:00 – 12:00 - Обзорная экскурсия по Тюмени')
    #     runner.bold = True
    #     runner.add_break()
    #     runner = day1_desc.add_run('''
    # ''')
    #     runner.add_break()
    for event in events:
        insert_event(doc, event['header'], event['body'])

    doc.save('demo.docx')


if __name__ == '__main__':
    main()
